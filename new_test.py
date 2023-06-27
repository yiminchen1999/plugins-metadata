import openai
import warnings
warnings.filterwarnings("ignore", category=UserWarning)

openai.api_key = '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'
import PyPDF2
import openpyxl
import pandas as pd
from docx import Document

# Read the Excel file
df = pd.read_excel('/Users/chenyimin/PycharmProjects/plugins-quickstart/CSCL_1995.xlsx', engine='openpyxl')

# Create a new column 'content' to store the paper content
df['content'] = ""

# Iterate through each row of the DataFrame
for index, row in df.iterrows():
    title = row['title']  # Assuming 'title' is the column name for the paper title

    # Open the Word document
    doc = Document('/Users/chenyimin/PycharmProjects/plugins-quickstart/1995.docx')

    # Search for the paper by its title
    found = False
    include_content = True  # Flag to include content (including abstract) after finding the paper
    for para in doc.paragraphs:
        if title in para.text:
            found = True
            include_content = True  # Start including content after finding the paper
        elif found and include_content:
            row['content'] += para.text



    # Iterate through tables in the document
    for table in doc.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                if title in cell.text:
                    found = True
                    include_content = True  # Start including content after finding the paper
                elif found and include_content:
                    row['content'] += cell.text



# Save the updated DataFrame to a new Excel file
df.to_excel('/Users/chenyimin/PycharmProjects/plugins-quickstart/1995output_ver2.xlsx', index=False)

